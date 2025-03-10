import requests,openpyxl,time
from docx import Document

singer_or_song=input('请输入歌手或歌曲名字')
numbers=int(input('请输入要爬取页数（每页最多爬取60首）'))
output=int(input('保存至excel请按1，保存至word请按2'))

for lines in list(singer_or_song):

    for i in range(0,numbers):
        url='https://c.y.qq.com/soso/fcgi-bin/client_search_cp'

        headers={
            'origin': 'https://y.qq.com',
            'referer': 'https://y.qq.com/',
            'user-agent': 'Mozilla/5.0 (Linux; Android 6.0; Nexus 5 Build/MRA58N) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/104.0.0.0 Mobile Safari/537.36'
            }
        

        params={
        '﻿ct':'24',
        'qqmusic_ver':'1298',
        'remoteplace':'txt.yqq.top',
        'searchid':'96381501447179318',
        'aggr':'0',
        'catZhida':'1',
        'lossless':'0',
        'sem':'1',
        't':'7',
        'p':str(i+1),
        'n':'5',
        'w':lines,
        '_':'1631713682489',
        'cv':'4747474',
        'ct':'24',
        'format':'json',
        'inCharset':'utf-8',
        'outCharset':'utf-8',
        'notice':'0',
        'platform':'yqq.json',
        'needNewCode':'0',
        'uin':'0',
        'g_tk_new_20200303':'5381',
        'g_tk':'5381',
        'hostUin':'0',
        'loginUin':'0'
        }

        res=requests.get(url,headers=headers,params=params)
        print(res)

        dict_lyrics=res.json()
        


        lyric_list=dict_lyrics['data']['lyric']['list']

        #i='\n '
        #list_section1=[]

        #list_combination=[]
        list_s1s2_onebyone=[]
        for item in lyric_list:
            
            lyric=item['content']
            list_split= lyric.split('\\n ')  
            list_section2=[]
            for i in range(1,len(list_split)):
                list_section1=[list_split[0]]#歌曲名
                list_section2.append(list_split[i])#作者和歌词
        
            list_s1s2=[list_section1,list_section2]
            list_s1s2_onebyone.append(list_s1s2)
        #print(list_section2)
        #print(list_s1s2_onebyone)
        
        list_for_sheet=[]
        for comb in list_s1s2_onebyone:
            #for row in comb:
            str_=''
            for r in range(0,len(comb[1])):
                str_=str_+comb[1][r]+'\n'#歌词主体部分
            #list_str=[str_]
            #print(comb[0][0])
            list_namesrt=[comb[0][0],str_]
            #print(list_namesrt)
            list_for_sheet.append(list_namesrt)

    time=time.asctime(time.localtime(time.time()) )

    if output==1:
        wb=openpyxl.Workbook()
        sheet=wb.active
        sheet.title='lyrics'
        head=['歌名','歌词']
        sheet.append(head)
        for lyric in list_for_sheet:
            sheet.append(lyric)
        sheet.append([''])
        sheet.append([time])
        wb.save('%s.xlsx'%(singer_or_song))

    elif output==2:
        document=Document()
        for lyric in list_for_sheet:
            for sentence in lyric:
                paragraph=document.add_paragraph(sentence)
            paragraph=document.add_paragraph(['\n'])
        document.save('%s.docx'%(singer_or_song))



#查看晴天歌词
#f=openpyxl.load_workbook('JayChowLyrics.xlsx')['lyrics']['B2'].value
#print(f)'''





    
